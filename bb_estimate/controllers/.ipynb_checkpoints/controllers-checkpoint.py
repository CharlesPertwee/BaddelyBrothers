# -*- coding: utf-8 -*-

from odoo import http, tools, _
from odoo.http import request, Controller
import requests
import codecs
import datetime
import base64
from docx import Document
from docx.shared import Cm, Mm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import io

class BbEstimate(http.Controller):
    @http.route('/bb_estimate/bb_estimate/estimateLetter/<string:id>', auth='public')
    def estimate_letter(self, **kw):
        values = dict(kw)
        Estimate = request.env['bb_estimate.estimate'].sudo().search([('id','=',values['id'])])
        document = Document()
        
        #url = 'https://charlespertwee-baddelybrothers-customization-dv-444705.dev.odoo.com/bb_estimate/static/src/img/BaddeleyNew.jpg'
        url = request.httprequest.host_url+"bb_estimate/static/src/img/BaddeleyNew.jpg"
        response = requests.get(url, stream=True)
        header = io.BytesIO(response.content)
        section = document.sections[0]   # Create a section
        sec_header = section.header   # Create header 
        header_tp = sec_header.add_paragraph(style='No Spacing')  # Add a paragraph in the header, you can add any anything in the paragraph
        header_run = header_tp.add_run()   # Add a run in the paragraph. In the run you can set the values 
        header_run.add_picture(header,width=Inches(7.26772)) 
        
        universalTableStyle = "borderColor:white"
        styles = document.styles['No Spacing']
        font = styles.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        
        table = document.add_table(6, 3)
        table.style = 'Table Grid'
        tbl = table._tbl # get xml element in table
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
            
        table.line_spacing_rule = 0
        table.cell(0,0).text = Estimate.partner_id.name
        table.cell(0,2).text = "Date: %s"%(str(datetime.datetime.now().strftime('%d-%m-%Y')))
        
        x = 1
        
        if Estimate.invoice_account.street:
            table.cell(x,0).text = Estimate.invoice_account.street
            table.cell(x,1).text = " "
            x += 1
            
        if Estimate.invoice_account.street2 :
            table.cell(x,0).text = Estimate.invoice_account.street2 
            table.cell(x,1).text = " "
            x += 1
            
        if Estimate.invoice_account.city:
            table.cell(x,0).text = Estimate.invoice_account.city 
            table.cell(x,1).text = " "
            x += 1
        
        if Estimate.invoice_account.zip:
            table.cell(x,0).text = Estimate.invoice_account.zip 
            table.cell(x,1).text = " "
            x += 1
          
        table.cell(x,0).text = "Attn: %s"%(Estimate.contact.name)
        table.cell(x,2).text = "Estimate Number: %s"%(str(Estimate.estimate_number))
        
        document.add_paragraph("Dear %s,"%(Estimate.contact.name.split()[0]))
        document.add_paragraph("Thank you for your enquiry, we have pleasure in submitting our estimate as follows:")
        
        record_length = len([x for x in Estimate.estimate_line if not x.isExtra])
        data_table = document.add_table(record_length+7, 3) 
        data_table.style = 'Table Grid'
        tblData = data_table._tbl # get xml element in table
        for cell in tblData.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
            
        data_table.cell(0,0).text = "Title: "      
        data_table.cell(0,1).text = Estimate.title
        
        data_table.cell(1,0).text = "Size: "      
        data_table.cell(1,1).text = "%s X %s"%(Estimate.finished_height,Estimate.finished_width)
        
        y = 2
        
        for estimate_line_material in Estimate.estimate_line:
            if estimate_line_material.option_type == 'material' and estimate_line_material.customer_description and not estimate_line_material.isExtra:
                data_table.cell(y,0).text = "Material: "
                data_table.cell(y,1).text = estimate_line_material.customer_description
                y += 1
        
        document.add_paragraph(Estimate.GenerateEnvelopeDetails(Estimate))
        
        for estimate_line_process in Estimate.estimate_line:
            if estimate_line_process.option_type == 'process' and estimate_line_process.customer_description and not estimate_line_process.isExtra:
                data_table.cell(y,0).text = "Process: "
                data_table.cell(y,1).text = estimate_line_process.customer_description
                y += 1
        
        if Estimate.quantity_1:
            data_table.cell(y,0).text="Qty/Price"
            data_table.cell(y,1).text="%s @ £%s"%(str(Estimate.quantity_1).encode("utf-8").decode("utf-8"),str(Estimate.total_price_1).encode("utf-8").decode("utf-8"))
        if Estimate.quantity_2:
            data_table.cell(y+1,0).text=""
            data_table.cell(y+1,1).text="%s @ £%s"%(str(Estimate.quantity_2).encode("utf-8").decode("utf-8"),str(Estimate.total_price_2).encode("utf-8").decode("utf-8"))
        if Estimate.quantity_3:
            data_table.cell(y+2,0).text=""
            data_table.cell(y+2,1).text="%s @ £%s"%(str(Estimate.quantity_3).encode("utf-8").decode("utf-8"),str(Estimate.total_price_3).encode("utf-8").decode("utf-8"))
        if Estimate.quantity_4:
            data_table.cell(y+3,0).text=""
            data_table.cell(y+3,1).text="%s @ £%s"%(str(Estimate.quantity_4).encode("utf-8").decode("utf-8"),str(Estimate.total_price_4).encode("utf-8").decode("utf-8"))
        if Estimate.total_price_run_on:
            data_table.cell(y+4,0).text=""
            data_table.cell(y+4,1).text="Run on: £%s per %s"%(str(Estimate.total_price_run_on).encode("utf-8").decode("utf-8"),str(Estimate.run_on).encode("utf-8").decode("utf-8"))
        
        for condition in Estimate.estimateConditions:
            document.add_paragraph(condition.description)
        
        if Estimate.hasExtra:
            extra_length = (len([x for x in Estimate.estimate_line if x.isExtra]))*6
            extra_table = document.add_table(extra_length, 3)            
            extra_table.style = 'Table Grid'
            tblExtraData = extra_table._tbl # get xml element in table
            for cell in tblExtraData.iter_tcs():
                tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'nil')

                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'nil')

                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'nil')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')

                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'nil')

                tcBorders.append(top)
                tcBorders.append(left)
                tcBorders.append(bottom)
                tcBorders.append(right)
                tcPr.append(tcBorders)
                
            z = 0
            for extra in Estimate.estimate_line:
                if extra.isExtra and extra.extraDescription:
                    extra_table.cell(z,0).text = "Extras"
                    extra_table.cell(z,1).text = extra.extraDescription
                    extra_table.cell(z+1,1).text = "%s @ £%s"%(str(extra.quantity_1).encode("utf-8").decode("utf-8"),str(extra.total_price_1).encode("utf-8").decode("utf-8"))
                    extra_table.cell(z+2,1).text = "%s @ £%s"%(str(extra.quantity_2).encode("utf-8").decode("utf-8"),str(extra.total_price_2).encode("utf-8").decode("utf-8"))
                    extra_table.cell(z+3,1).text = "%s @ £%s"%(str(extra.quantity_3).encode("utf-8").decode("utf-8"),str(extra.total_price_3).encode("utf-8").decode("utf-8"))
                    extra_table.cell(z+4,1).text = "%s @ £%s"%(str(extra.quantity_4).encode("utf-8").decode("utf-8"),str(extra.total_price_4).encode("utf-8").decode("utf-8"))
                    extra_table.cell(z+5,1).text = "Run on: £%s per %s"%(str(extra.total_price_run_on).encode("utf-8").decode("utf-8"),str(extra.run_on).encode("utf-8").decode("utf-8"))
                    z += 6
                           
        document.add_paragraph("Should you require any more information or wish to discuss your detailed requirements further, please contact us on 020 8986 2666.")
        
        document.add_paragraph('Yours faithfully,')
        document.add_paragraph(Estimate.estimator.name)
        
        #url1 = 'https://charlespertwee-baddelybrothers-customization-dv-444705.dev.odoo.com/bb_estimate/static/src/img/BaddeleyFooter.png'
        url1 = request.httprequest.host_url+"bb_estimate/static/src/img/BaddeleyFooter.png"
        response1 = requests.get(url1, stream=True)
        footer = io.BytesIO(response1.content)
        section = document.sections[0]
        default_footer = section.footer   # Add a footer
        footer_p = default_footer.add_paragraph()
        footer_r = footer_p.add_run()
        footer_r.add_picture(footer,width=Inches(6.26772)) 
        
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()  
        
        attachment = request.env['ir.attachment'].create({
            'name': "%s.docx"%(Estimate.estimate_number),
            'type':'binary',
            'res_model':'bb_estimate.estimate',
            'res_id':Estimate.id,
            'datas_fname':"%s.doc"%(Estimate.estimate_number),
            'mimetype':'application/msword',
            'datas':base64.encodestring(docx_bytes)
        })
        
        pdfhttpheaders = [('Content-Type','application/msword'),("Content-Disposition","filename= %s.docx"%(Estimate.estimate_number))]       
        return request.make_response(docx_bytes, headers=pdfhttpheaders)
       
        
