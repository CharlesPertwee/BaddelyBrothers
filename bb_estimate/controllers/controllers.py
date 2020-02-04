# -*- coding: utf-8 -*-

from odoo import http, tools, _
from odoo.http import request, Controller
import requests
import codecs
import datetime
import base64
from docx import Document
from docx.shared import Cm, Mm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import io
import traceback
import os

DIE_SIZES = [
    ('none', ''),
    ('standard','No Die (No Charge)'),
    ('small','Crest Die'),
    ('medium','Heading Die'),
    ('large','Invitation Die')
]

ENVELOPE_TYPES = [
    ('none', ''),
    ('diamond', 'Diamond'),
    ('highcutdiamond','High Cut Diamond'),
    ('tuck', 'Tuck'),
    ('pocket', 'Pocket'),
    ('walletpocket','Walletstyle Pocket'),
    ('wallet', 'Wallet'),
    ('banker', 'Banker'),
    ('tissuelined', 'Tissue Lined, Embossed & Windowed'),
]

FLAP_GLUE_TYPES = [
    ('none', ''),
    ('gummed', 'Gummed'),
    ('doublegummed','Double Gummed'),
    ('peel', 'Peel & Stick'),
    ('ungummed', 'Un-gummed'),
    ('topless', 'Topless'),
    ('stringwasher', 'String & Washer'),
]

TISSUE_LINING_OPTIONS = [
    ('none', ''),
    ('full', 'Yes - Fully'),
    ('half', 'Yes - Half'),
    ('unlined', 'Unlined'),
]

class BbEstimate(http.Controller):
    @http.route('/bb_estimate/bb_estimate/estimateLetter/<string:id>', auth='public')
    def estimate_letter(self, **kw):
        values = dict(kw)
        Estimate = request.env['bb_estimate.estimate'].sudo().search([('id','=',values['id'])])
        document = Document()
        paragraph_format = document.styles['No Spacing']
        section = document.sections[0]   # Create a section
        section.header_distance = Cm(0)# ADD HEADER DISTANCE
        section.left_margin = Mm(24)# ADD LEFT MARGIIN
        section.right_margin = Mm(25)# ADD RIGHT MARGIN
        sec_header = section.header   # Create header 
        sec_header.left_margin = Cm(0)
        header_tp = sec_header.add_paragraph(style='No Spacing')  # Add a paragraph in the header, you can add any anything in the paragraph
        header_run = header_tp.add_run()   # Add a run in the paragraph. In the run you can set the values 
        header_run.add_picture('/opt/Bb/BaddelyBrothers/bb_estimate/static/src/img/HeaderNew.jpg',width=Inches(6.57)) 
        universalTableStyle = "borderColor:white"
        styles = document.styles['No Spacing']
        font = styles.font
        font.name = 'Tahoma'
        font.size = Pt(12)
        
        
        table = document.add_table(6, 3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        tbl = table._tbl # get xml element in table
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
            
        table.line_spacing_rule = 0
        table.autofit = True
        table.cell(0,0).text = Estimate.partner_id.name
        table.cell(0,2).paragraphs[0].text = "Date: %s"%(str(datetime.datetime.now().strftime('%d/%m/%Y')))
        table.cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        #.add_paragraph("Date: %s"%(str(datetime.datetime.now().strftime('%d/%m/%Y'))))
        
        #table.cell(0,2).vertical_alignment = WD_ALIGN_VERTICAL.TOP
        #paraDate.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        
        x = 1
        
        if Estimate.invoice_account.street:
            table.cell(x,0).text = Estimate.invoice_account.street
            table.cell(x,1).text = " "
            x += 1
            
        if Estimate.invoice_account.street2 :
            table.cell(x,0).text = Estimate.invoice_account.street2 
            table.cell(x,1).text = " "
            x += 1
            
        if Estimate.invoice_account.city:
            table.cell(x,0).text = Estimate.invoice_account.city 
            table.cell(x,1).text = " "
            x += 1
        
        if Estimate.invoice_account.zip:
            table.cell(x,0).text = Estimate.invoice_account.zip 
            table.cell(x,1).text = " "
            x += 1
          
        table.cell(x,0).text = "\nAttn: %s"%(Estimate.contact.name)
        
        table.cell(x,2).paragraphs[0].text = "\nEstimate no: %s"%(str(Estimate.estimate_number))
        table.cell(x,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
#         para = table.cell(x,2).add_paragraph("Estimate no: %s"%(str(Estimate.estimate_number)))
#         para.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        
#         table.cell(x,2).paragraphs[0].text = "Estimate no: %s"%(str(Estimate.estimate_number))
#         table.cell(x,2).paragraphs[0].text = WD_ALIGN_PARAGRAPH.RIGHT
        
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name= "Tahoma"
        
        parag1 = document.add_paragraph("")
        paragraph_format1 = parag1.paragraph_format
        paragraph_format1.left_indent = Inches(0.09)
        senten1 = parag1.add_run("\nDear %s,"%(Estimate.contact.name.split()[0]))
        senten1.font.name = "Tahoma"
        senten1 = parag1.add_run("\nThank you for your enquiry, we have pleasure in submitting our estimate as follows:")
        senten1.font.name = "Tahoma"

        new_col = 2
        if Estimate.isEnvelope:
            if (Estimate.envelope_type) and (not Estimate.envelope_type == 'none'):
                new_col += 1

            if (Estimate.flap_glue_type) and (not Estimate.flap_glue_type == 'none'):
                new_col += 1

            if (Estimate.tissue_lined) and (not Estimate.tissue_lined == 'none'):
                new_col += 1

            if Estimate.embossed:
                new_col += 1

            if Estimate.windowed:
                if not Estimate.standardWindowSize:
                    new_col += 1
                else:
                    new_col += 2
        
        record_length = len([x for x in Estimate.estimate_line if x.customer_description and (not x.isExtra)])
        addtitional = sum([1 for x in ['1','2','3','4','run_on'] if abs(Estimate['total_price_'+x] - Estimate['total_price_extra_'+x]) >0]) + new_col
        data_table = document.add_table((record_length + addtitional) , 2)  
        # raise Exception((record_length + addtitional))
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        data_table.allow_autofit = True
        data_table.autofit = True
        data_table.style = 'Table Grid'
        tblData = data_table._tbl # get xml element in table
        for cell in tblData.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
        
        inches = 1.5
        
        data_table.cell(0,0).text = "Title: "   
        #data_table[0].width = Inches(inches)
        data_table.cell(0,1).text = Estimate.title
        
        data_table.cell(1,0).text = "Size: "    
        #data_table[0].width = Inches(inches)
        data_table.cell(1,1).text = "%s X %s\n"%(Estimate.finished_height,Estimate.finished_width)
        data_table.cell(1,1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
       
        y = 2             

        for estimate_line_material in Estimate.estimate_line:
            if estimate_line_material.option_type == 'material' and estimate_line_material.customer_description and estimate_line_material.documentCatergory and (not estimate_line_material.isExtra) and ( not estimate_line_material.material.productType=='Delivery'):
                data_table.cell(y,0).text = "%s: "%(estimate_line_material.documentCatergory)
                #data_table.cell(y,0).width = Inches(inches)
                data_table.cell(y,1).text = estimate_line_material.customer_description
                y += 1

        if Estimate.isEnvelope:
            if (Estimate.envelope_type) and (not Estimate.envelope_type == 'none'):
                data_table.cell(y,0).text = "Type: "   
                data_table.cell(y,1).text = dict(ENVELOPE_TYPES)[Estimate.envelope_type]
                y += 1

            if (Estimate.flap_glue_type) and (not Estimate.flap_glue_type == 'none'):
                data_table.cell(y,0).text = "Flap: "   
                data_table.cell(y,1).text = dict(FLAP_GLUE_TYPES)[Estimate.flap_glue_type]
                y += 1

            if (Estimate.tissue_lined) and (not Estimate.tissue_lined == 'none'):
                data_table.cell(y,0).text = "Lining: "   
                data_table.cell(y,1).text = dict(TISSUE_LINING_OPTIONS)[Estimate.tissue_lined]
                y += 1

            if Estimate.embossed:
                data_table.cell(y,0).text = "Process: "   
                data_table.cell(y,1).text ="Blind Embossed"
                y += 1

            if Estimate.windowed:
                if not Estimate.standardWindowSize:
                    data_table.cell(y,0).text = "Window Size: "   
                    data_table.cell(y,1).text ="Standard"
                    y += 1

                else:
                    data_table.cell(y,0).text = "Window Size: "
                    data_table.cell(y,1).text = '%s mm  x  %s mm' % (Estimate.windowHeight, Estimate.windowWidth)
                    y += 1
                    data_table.cell(y,0).text = "Window Pos: "
                    data_table.cell(y,1).text = '%s mm FLHS,  %s mm Up' % (Estimate.windowFlhs, Estimate.windowUp)
                    y += 1
        
        
        for estimate_line_process in Estimate.estimate_line:
            if estimate_line_process.option_type == 'process' and estimate_line_process.documentCatergory and estimate_line_process.customer_description and (not estimate_line_process.isExtra):
                data_table.cell(y,0).text = "%s: "%(estimate_line_process.documentCatergory)
                #data_table.cell(y,0).width = Inches(inches)
                data_table.cell(y,1).text = estimate_line_process.customer_description
                y += 1
        for estimate_line_material in Estimate.estimate_line:   
            if estimate_line_material.option_type == 'material' and estimate_line_material.documentCatergory and estimate_line_material.customer_description and (not estimate_line_material.isExtra) and estimate_line_material.material.productType=='Delivery':
                data_table.cell(y,0).text = "%s: "%(estimate_line_material.documentCatergory)
                #data_table.cell(y,0).width = Inches(inches)
                data_table.cell(y,1).text = estimate_line_material.customer_description
                y += 1
        
        if abs(Estimate.total_price_1 - Estimate.total_price_extra_1) >0:
            data_table.cell(y,0).text="Qty/Price:"
            price = 0
            if Estimate.total_price_1 > Estimate.total_price_extra_1:
                price =  (Estimate.total_price_1 - Estimate.total_price_extra_1)
            else:
                price = (Estimate.total_price_extra_1 - Estimate.total_price_1)
                
            if price:                
                data_table.cell(y,1).text="%d    @    £%.2f"%(Estimate.quantity_1,price)
                y+=1
        
        if abs(Estimate.total_price_2 - Estimate.total_price_extra_2) >0:
            data_table.cell(y,0).text=""
            price = 0
            if Estimate.total_price_2 > Estimate.total_price_extra_2:
                price =  (Estimate.total_price_2 - Estimate.total_price_extra_2)
            else:
                price = (Estimate.total_price_extra_2 - Estimate.total_price_2)
                
            if price:                
                data_table.cell(y,1).text="%d    @    £%.2f"%(Estimate.quantity_2,price)
                y+=1
        
        if abs(Estimate.total_price_3 - Estimate.total_price_extra_3) >0:
            data_table.cell(y,0).text=""
            price = 0
            if Estimate.total_price_3 > Estimate.total_price_extra_3:
                price =  (Estimate.total_price_3 - Estimate.total_price_extra_3)
            else:
                price = (Estimate.total_price_extra_3 - Estimate.total_price_3)
            
            if price:            
                data_table.cell(y,1).text="%s    @    £%.2f"%(Estimate.quantity_3,price)
                y+=1
        
        if abs(Estimate.total_price_4 - Estimate.total_price_extra_4) >0:
            data_table.cell(y,0).text=""
            price = 0
            if Estimate.total_price_4 > Estimate.total_price_extra_4:
                price =  (Estimate.total_price_4 - Estimate.total_price_extra_4)
            else:
                price = (Estimate.total_price_extra_4 - Estimate.total_price_4)
            
            if price:                   
                data_table.cell(y,1).text="%s    @    £%.2f"%(Estimate.quantity_4,price)
                y+=1
        
        if abs(Estimate.total_price_run_on - Estimate.total_price_extra_run_on) >0:
            data_table.cell(y,0).text=""
            price = 0
            if Estimate.total_price_run_on > Estimate.total_price_extra_run_on:
                price =  (Estimate.total_price_run_on - Estimate.total_price_extra_run_on)
            else:
                price = (Estimate.total_price_extra_run_on - Estimate.total_price_run_on)
            if price:
                data_table.cell(y,1).text="Run on: £%.2f per %s"%(price,Estimate.run_on)
                y+=1
        
        
        
        for row in data_table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                row.cells[0].width = Inches(inches)
                row.cells[1].width = Inches(5.06)
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name= "Tahoma"
                        
        #for column in data_table.columns[0]:
        data_table.columns[0].width = Inches(inches)
        data_table.columns[1].width = Inches(5.06)
            
#         for row in data_table.rows:
#             row.cells[0].width = Inches(0.5)
            
        #run = document.add_paragraph().add_run()
        #run.add_break() 
        
        line = ""
        line = document.add_paragraph("")
#         run = line.add_run()
#         run.add_break() 

        # envelopeLine = ""
        # if Estimate.isEnvelope:
        #     envelopeLine = document.add_paragraph(Estimate.GenerateEnvelopeDetails(Estimate),style=paragraph_format)
        #     paragraph_format4 = envelopeLine.paragraph_format
        #     paragraph_format4.left_indent = Inches(0.09)
        #     envelopeLine.style.font.size = Pt(8)
            
        # if envelopeLine:
        #     runEnv = envelopeLine.add_run()
        #     runEnv.add_break()
        
        if Estimate.hasExtra:
            extra_length = (len([x for x in Estimate.estimate_line if (x.isExtra and x.extraDescription)]))
            addtitional1 = sum([1 for x in ['total_price_extra_1','total_price_extra_2','total_price_extra_3','total_price_extra_4','total_price_extra_run_on'] if Estimate[x] > 0]) + new_col
            extra_table = document.add_table(extra_length*(addtitional1-new_col+1), 2)  
            extra_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            extra_table.style = 'Table Grid'
            tblExtraData = extra_table._tbl # get xml element in table
            for cell in tblExtraData.iter_tcs():
                tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'nil')

                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'nil')

                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'nil')
                bottom.set(qn('w:sz'), '4')
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')

                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'nil')

                tcBorders.append(top)
                tcBorders.append(left)
                tcBorders.append(bottom)
                tcBorders.append(right)
                tcPr.append(tcBorders)
                
            z = 0
            for extra in Estimate.estimate_line:
                if extra.isExtra and extra.extraDescription:
                    extra_table.cell(z,0).text = "Extras:"
                    extra_table.cell(z,1).text = extra.extraDescription
                    if extra.quantity_1 and extra.total_price_1:
                        z += 1
                        extra_table.cell(z,1).text = "%s    @    £%s"%(str(extra.quantity_1).encode("utf-8").decode("utf-8"),str(extra.total_price_1).encode("utf-8").decode("utf-8"))
                    if extra.quantity_2 and extra.total_price_2:
                        z += 1
                        extra_table.cell(z,1).text = "%s    @    £%s"%(str(extra.quantity_2).encode("utf-8").decode("utf-8"),str(extra.total_price_2).encode("utf-8").decode("utf-8"))
                    if extra.quantity_3 and extra.total_price_3:
                        z += 1
                        extra_table.cell(z,1).text = "%s    @    £%s"%(str(extra.quantity_3).encode("utf-8").decode("utf-8"),str(extra.total_price_3).encode("utf-8").decode("utf-8"))
                    if extra.quantity_4 and extra.total_price_4:
                        z += 1
                        extra_table.cell(z,1).text = "%s    @    £%s"%(str(extra.quantity_4).encode("utf-8").decode("utf-8"),str(extra.total_price_4).encode("utf-8").decode("utf-8"))
                    if extra.run_on and extra.total_price_run_on:
                        z += 1
                        extra_table.cell(z,1).text = "Run on: £ %s per %s"%(str(extra.total_price_run_on).encode("utf-8").decode("utf-8"),str(extra.run_on).encode("utf-8").decode("utf-8"))
                    z += 1
                    
            for row in extra_table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    row.cells[0].width = Inches(inches)
                    row.cells[1].width = Inches(5.06)
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.name= "Tahoma"
                            
            extra_table.columns[0].width = Inches(inches)
            extra_table.columns[1].width = Inches(5.06)
            
        
        
        
        line = ""
        spaceApplied = False
        for condition in Estimate.estimateConditions:
            line = document.add_paragraph(condition.description,style=paragraph_format)
            paragraph_format2 = line.paragraph_format
            if not spaceApplied:
                paragraph_format2.space_before = Pt(12)
                spaceApplied = True
            paragraph_format2.left_indent = Inches(0.09)

            line.style.font.size = Pt(9)
        
        if line:
            run = line.add_run()
            run.add_break() 
        
        p = document.add_paragraph("")
        paragraph_format3 = p.paragraph_format
        paragraph_format3.left_indent = Inches(0.09)
        sen = p.add_run('\n\nShould you require any more information or wish to discuss your detailed requirements further, please contact us on 020 8986 2666.')
        sen.font.name = 'Tahoma'
        
        p1 = document.add_paragraph('')
        paragraph_format3 = p1.paragraph_format
        paragraph_format3.left_indent = Inches(0.09)
        sen1 = p1.add_run('Yours sincerely,')
        sen1.font.name = 'Tahoma'
        sen1.add_break()
        sen2 = p1.add_run(Estimate.estimator.name)
        sen2.font.name = 'Tahoma'
        

        section = document.sections[0]
        default_footer = section.footer   # Add a footer
        footer_p = default_footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_r = footer_p.add_run()
        footer_r.add_picture('/opt/Bb/BaddelyBrothers/bb_estimate/static/src/img/FooterNew.jpg',width=Inches(3.5)) 
        
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()  
        
        attachment = request.env['ir.attachment'].create({
            'name': "%s.docx"%(Estimate.estimate_number),
            'type':'binary',
            'res_model':'bb_estimate.estimate',
            'res_id':Estimate.id,
            'datas_fname':"%s.docx"%(Estimate.estimate_number),
            'mimetype':'application/msword',
            'datas':base64.encodestring(docx_bytes)
        })
        pdfhttpheaders = [('Content-Type','application/msword'),("Content-Disposition","filename= %s.docx"%(Estimate.estimate_number))]       
        return request.make_response(docx_bytes, headers=pdfhttpheaders)
       
        
        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)
