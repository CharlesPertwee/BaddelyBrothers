# -*- coding: utf-8 -*-
from odoo import http, tools, _
from odoo.http import request, Controller
from docx import Document
from docx.shared import Mm, Inches, Pt
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class BbContacts(http.Controller):
     @http.route(['/doc/report/<model("bb_contacts.contacts_link"):id>/<string:type>'], type='http', auth='public')
     def report_docx(self,**kw):
        model_id = kw.get('id',"NONE").id
        type = kw.get('type','None')
        model = request.env['bb_contacts.contacts_link'].sudo().search([('id','=',model_id)])
        document = Document()
        if type == "normal":
            section = document.sections[0]
            section.page_height = Mm(100)
            section.page_width = Mm(104)

            style = document.styles['Normal']
            font = style.font
            font.name = 'Tahoma'
            font.size = Pt(12)

            p = document.add_paragraph(model.contact.name)
            paragraph_format = p.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p = document.add_paragraph(model.company.name)
            paragraph_format = p.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT           
        
        if type == 'letter':
            raise Exception('SPARTA')
            contact = document.add_paragraph(model.contact.name)
            paragraph_format_contact = contact.paragraph_format
            paragraph_format_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            company = document.add_paragraph(model.company.name)
            paragraph_format_company = company.paragraph_format
            paragraph_format_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            date = document.add_paragraph(datetime.today().date())
            paragraph_date = date.paragraph_format
            paragraph_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            
         
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()            

        pdfhttpheaders = [('Content-Type','application/msword')]        
        return request.make_response(docx_bytes, headers=pdfhttpheaders)
            
            
            
            
            


            #p = document.add_paragraph('A plain paragraph having some ')
            #p.add_run('bold').bold = True
            #p.add_run(' and some ')
            #p.add_run('italic.').italic = True

            #document.add_heading('Heading, level 1', level=1)
            #document.add_paragraph('Intense quote', style='Intense Quote')

            #document.add_paragraph(
            #    'first item in unordered list', style='List Bullet'
            #)
            #document.add_paragraph(
            #    'first item in ordered list', style='List Number'
            #)

            #records = (
            #    (3, '101', 'Spam'),
            #    (7, '422', 'Eggs'),
            #    (4, '631', 'Spam, spam, eggs, and spam')
            #)

            #table = document.add_table(rows=1, cols=3)
            #hdr_cells = table.rows[0].cells
            #hdr_cells[0].text = 'Qty'
            #hdr_cells[1].text = 'Id'
            #hdr_cells[2].text = 'Desc'
            #for qty, id, desc in records:
            #    row_cells = table.add_row().cells
            #    row_cells[0].text = str(qty)
            #    row_cells[1].text = id
            #    row_cells[2].text = desc
    #
            #document.add_page_break()
