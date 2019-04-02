# -*- coding: utf-8 -*-
from odoo import http, tools, _
from odoo.http import request, Controller
import requests
from docx import Document
from docx.shared import Mm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import io

class BbContacts(http.Controller):
     @http.route(['/doc/report/<model("bb_contacts.contacts_link"):id>/<string:type>/<string:date>'], type='http', auth='public')
     def report_docx(self,**kw):
        model_id = kw.get('id',"NONE").id
        type = kw.get('type','None')
        date =  kw.get('date','None')
        model = request.env['bb_contacts.contacts_link'].sudo().search([('id','=',model_id)])
        document = Document()
        
        style = document.styles['Normal']
        font = style.font
        font.name = 'Tahoma'
        font.size = Pt(12)
        
        if type == "normal":
            section = document.sections[0]
            section.page_height = Mm(100)
            section.page_width = Mm(104)

            p = document.add_paragraph(model.contact.name)
            paragraph_format = p.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            p = document.add_paragraph(model.company.name)
            paragraph_format = p.paragraph_format
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT           
        
        if type == 'letter' or type == 'head_letter':
            #if type == 'head_letter':
            #   url = 'https://upload.wikimedia.org/wikipedia/commons/0/0e/Logo12345.png'
            #    response = requests.get(url, stream=True)
            #    image = io.BytesIO(response.content)
            #    document.add_picture(image, width=Inches(1.25))              
                                
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            contact = document.add_paragraph(model.contact.name)
            paragraph_format_contact = contact.paragraph_format
            paragraph_format_contact.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            company = document.add_paragraph(model.company.name)
            paragraph_format_company = company.paragraph_format
            paragraph_format_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            document.add_paragraph(" ")
            
            date = document.add_paragraph(date)
            paragraph_date = date.paragraph_format
            paragraph_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
            dear = document.add_paragraph("Dear,")
            paragraph_format_dear = dear.paragraph_format
            paragraph_format_dear.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            start = document.add_paragraph("Herewith enclosed, please find as requested:")
            paragraph_format_start = start.paragraph_format
            paragraph_format_start.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            document.add_paragraph(" ")
            
            document.add_paragraph("Should you require any more information or wish to discuss your detailed requirements further, please contact us on 020 8986 2666.")
            
            document.add_paragraph(" ")
            
            document.add_paragraph("Yours faithfully sincerely,")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph(" ")
            document.add_paragraph("Steve Wood")
            
         
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_bytes = docx_stream.getvalue()            

        pdfhttpheaders = [('Content-Type','application/msword')]        
        return request.make_response(docx_bytes, headers=pdfhttpheaders)


   
              
            
        